"""
Count images and check what figures exist in CSE493 Chapter 4.
Also extract images info from the docx.
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os

BASE = r"c:\Users\cmyli\OneDrive - eiu.edu.vn\Desktop\CSE492-main"
file_493 = os.path.join(BASE, "CSE493_MotoMarket_Motorcycle Marketplace and Service Hub_Pham Minh Tuan_Dang Pham Huu Thao.docx")

doc = Document(file_493)

# Count images in the document
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_count += 1

print(f"Total images in CSE493: {image_count}")

# Find all Caption paragraphs
print("\n" + "=" * 80)
print("ALL FIGURE CAPTIONS:")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ""
    if 'Caption' in style or (text and text.startswith('Figure')):
        print(f"  [{i}] {text}")

# Count images in Chapter 4 area (by looking at inline shapes near chapter 4 paragraphs)
print("\n" + "=" * 80)
print("INLINE SHAPES (images) in document:")
print("=" * 80)
inline_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            inline_count += 1
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
            inline_count += 1

print(f"Inline shapes found: {inline_count}")

# Also check for images within paragraphs using XML namespace
from lxml import etree
ns = {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
      'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
      'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
      'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

in_ch4 = False
ch4_images = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ""
    if 'toc' in style.lower():
        continue
    if 'Chapter 4' in text:
        in_ch4 = True
    if in_ch4 and 'Chapter 5' in text:
        break
    if in_ch4:
        drawings = para._element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if drawings:
            ch4_images += len(drawings)
            print(f"  Image found near para [{i}]: {text[:80] if text else '(empty text - image only)'}")

print(f"\nTotal images in Chapter 4 area: {ch4_images}")
