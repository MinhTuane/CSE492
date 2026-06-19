"""
Script to extract and compare text content from CSE492 and CSE493 docx reports.
"""
import sys
try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '--quiet'])
    from docx import Document

import os

BASE = r"c:\Users\cmyli\OneDrive - eiu.edu.vn\Desktop\CSE492-main"

def extract_headings_and_structure(filepath):
    """Extract headings and paragraph structure from a docx file."""
    doc = Document(filepath)
    result = []
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "None"
        text = para.text.strip()
        if not text:
            continue
        if 'Heading' in style_name or 'heading' in style_name:
            result.append(f"[{style_name}] {text}")
        elif 'Title' in style_name:
            result.append(f"[{style_name}] {text}")
    return result

def extract_full_text(filepath):
    """Extract all text from a docx file."""
    doc = Document(filepath)
    texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    return texts

# Extract headings from both files
file_492 = os.path.join(BASE, "CSE492_MotoMarket_Motorcycle Marketplace and Service Hub_Pham Minh Tuan_Dang Pham Huu Thao.docx")
file_493 = os.path.join(BASE, "CSE493_MotoMarket_Motorcycle Marketplace and Service Hub_Pham Minh Tuan_Dang Pham Huu Thao.docx")

print("=" * 80)
print("CSE492 HEADINGS STRUCTURE:")
print("=" * 80)
headings_492 = extract_headings_and_structure(file_492)
for h in headings_492:
    print(h)

print("\n" + "=" * 80)
print("CSE493 HEADINGS STRUCTURE:")
print("=" * 80)
headings_493 = extract_headings_and_structure(file_493)
for h in headings_493:
    print(h)

# Find differences in headings
print("\n" + "=" * 80)
print("HEADINGS IN CSE493 BUT NOT IN CSE492:")
print("=" * 80)
set_492 = set(headings_492)
set_493 = set(headings_493)
for h in headings_493:
    if h not in set_492:
        print(f"  + {h}")

print("\n" + "=" * 80)
print("HEADINGS IN CSE492 BUT NOT IN CSE493:")
print("=" * 80)
for h in headings_492:
    if h not in set_493:
        print(f"  - {h}")
